"""
通用文档转换器 —— 将多种格式的文档转为 Markdown。

支持的格式：
    PDF  → .md    (使用 pymupdf)
    EPUB → .md    (使用 ebooklib + BeautifulSoup)
    DOCX → .md    (使用 python-docx)
    HTML → .md    (使用 BeautifulSoup)

用法：
    from src.tools.doc_converter import convert_file, convert_directory
    convert_file("book.pdf")
    convert_directory("./my_books", "./output")
"""

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


# 支持的格式映射
SUPPORTED_FORMATS = {
    ".pdf":  "PDF 文档",
    ".epub": "EPUB 电子书",
    ".docx": "Word 文档",
    ".html": "HTML 网页",
    ".htm":  "HTML 网页",
}

# 最大文件大小限制 (100MB)
MAX_FILE_SIZE_MB = 100


@dataclass
class ConvertResult:
    """转换结果"""
    input_path: Path          # 输入文件路径
    output_path: Path         # 输出 .md 路径
    format: str               # 原始格式
    pages_or_chapters: int    # 页数/章节数
    size_kb: int              # 输出文件大小


# ============================================================
# 统一入口
# ============================================================

def convert_file(
    file_path: str | Path,
    output_dir: Optional[str | Path] = None,
) -> ConvertResult:
    """
    将单个文件转换为 Markdown，自动根据后缀选择转换器。

    Args:
        file_path:  输入文件路径
        output_dir: 输出目录（默认与源文件同目录）

    Returns:
        ConvertResult 对象

    Raises:
        ValueError: 不支持的格式
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_FORMATS:
        raise ValueError(f"不支持的格式: {suffix}（支持: {', '.join(SUPPORTED_FORMATS.keys())}）")

    out_dir = Path(output_dir) if output_dir else path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    if suffix == ".pdf":
        return _convert_pdf(path, out_dir)
    elif suffix == ".epub":
        return _convert_epub(path, out_dir)
    elif suffix == ".docx":
        return _convert_docx(path, out_dir)
    elif suffix in (".html", ".htm"):
        return _convert_html(path, out_dir)
    else:
        raise ValueError(f"不支持的格式: {suffix}")


def convert_directory(
    input_dir: str | Path,
    output_dir: Optional[str | Path] = None,
    recursive: bool = True,
) -> list[ConvertResult]:
    """
    批量转换目录下所有支持的文档。

    Args:
        input_dir:  输入目录
        output_dir: 输出目录（默认与输入目录相同）
        recursive:  是否递归子目录

    Returns:
        ConvertResult 列表
    """
    in_dir = Path(input_dir)
    out_dir = Path(output_dir) if output_dir else in_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    results: list[ConvertResult] = []
    extensions = tuple(SUPPORTED_FORMATS.keys())

    for ext in extensions:
        matches = in_dir.rglob(f"*{ext}") if recursive else in_dir.glob(f"*{ext}")
        for file_path in matches:
            # 跳过隐藏文件
            if file_path.name.startswith("."):
                continue
            # 跳过过大的文件
            if file_path.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
                continue
            try:
                result = convert_file(file_path, out_dir)
                results.append(result)
            except Exception as e:
                print(f"  ⚠️ 跳过 {file_path.name}: {e}")

    return results


# ============================================================
# PDF → Markdown
# ============================================================

def _convert_pdf(path: Path, out_dir: Path) -> ConvertResult:
    """使用 pymupdf 提取 PDF 文本并转为 Markdown"""
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    md_parts: list[str] = []
    page_count = doc.page_count

    # 标题
    title = path.stem
    md_parts.append(f"# {title}\n")
    md_parts.append(f"> 来源: {path.name} | 页数: {page_count}\n\n")

    for page_num in range(page_count):
        page = doc[page_num]
        text = page.get_text("text", sort=True)

        if text.strip():
            # 尝试检测章节标题（粗体、大字号）
            blocks = page.get_text("dict", sort=True).get("blocks", [])
            for block in blocks:
                if block["type"] == 0:  # 文本块
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            txt = span["text"].strip()
                            if not txt:
                                continue
                            # 粗体 + 短文本 → 可能是标题
                            font_size = span.get("size", 10)
                            is_bold = "Bold" in span.get("font", "")
                            if is_bold and len(txt) < 80 and font_size > 11:
                                md_parts.append(f"\n## {txt}\n\n")
                            else:
                                md_parts.append(txt + "\n")
                        md_parts.append("\n")
                else:
                    # 非文本块（图片等），跳过
                    pass
        else:
            # fallback: 直接用文本
            md_parts.append(text)

        # 每页添加分页标记
        if page_num < page_count - 1:
            md_parts.append(f"\n<!-- 第 {page_num + 1} 页 -->\n")

    doc.close()

    out_path = out_dir / f"{path.stem}.md"
    content = "".join(md_parts)
    _write_md(out_path, content)

    return ConvertResult(
        input_path=path,
        output_path=out_path,
        format="PDF",
        pages_or_chapters=page_count,
        size_kb=out_path.stat().st_size // 1024,
    )


# ============================================================
# EPUB → Markdown
# ============================================================

def _convert_epub(path: Path, out_dir: Path) -> ConvertResult:
    """使用 ebooklib 提取 EPUB 内容并转为 Markdown"""
    from bs4 import BeautifulSoup
    from ebooklib import epub

    book = epub.read_epub(str(path))
    md_parts: list[str] = []
    chapter_count = 0

    # 标题
    title = path.stem
    for item in book.get_metadata("DC", "title"):
        if item[0]:
            title = item[0]
            break

    md_parts.append(f"# {title}\n")
    md_parts.append(f"> 来源: {path.name}\n\n")

    # 读取文档内容
    for item in book.get_items_of_type(9):  # ITEM_DOCUMENT = 9
        content = item.get_content().decode("utf-8", errors="ignore")
        soup = BeautifulSoup(content, "lxml")

        # 提取标题
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
            level = int(tag.name[1])
            text = tag.get_text(strip=True)
            if text:
                md_parts.append(f"\n{'#' * level} {text}\n\n")

        # 提取段落
        for p in soup.find_all("p"):
            text = p.get_text(strip=True)
            if text and len(text) > 5:  # 过滤太短的
                md_parts.append(f"{text}\n\n")

        chapter_count += 1

    out_path = out_dir / f"{path.stem}.md"
    _write_md(out_path, "".join(md_parts))

    return ConvertResult(
        input_path=path,
        output_path=out_path,
        format="EPUB",
        pages_or_chapters=chapter_count,
        size_kb=out_path.stat().st_size // 1024,
    )


# ============================================================
# DOCX → Markdown
# ============================================================

def _convert_docx(path: Path, out_dir: Path) -> ConvertResult:
    """使用 python-docx 提取 Word 文档内容并转为 Markdown"""
    from docx import Document

    doc = Document(str(path))
    md_parts: list[str] = []
    para_count = len(doc.paragraphs)

    # 标题
    md_parts.append(f"# {path.stem}\n")
    md_parts.append(f"> 来源: {path.name} | 段落数: {para_count}\n\n")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_parts.append("\n")
            continue

        style = para.style.name if para.style else ""

        # 根据 Word 样式映射 Markdown 标题
        if style.startswith("Heading") or style.startswith("标题"):
            try:
                level = int(re.findall(r"\d+", style)[0])
            except (IndexError, ValueError):
                level = 2
            level = min(level, 6)
            md_parts.append(f"\n{'#' * level} {text}\n\n")
        elif style == "List Paragraph" or style == "列表":
            md_parts.append(f"- {text}\n")
        else:
            md_parts.append(f"{text}\n\n")

    out_path = out_dir / f"{path.stem}.md"
    _write_md(out_path, "".join(md_parts))

    return ConvertResult(
        input_path=path,
        output_path=out_path,
        format="DOCX",
        pages_or_chapters=para_count,
        size_kb=out_path.stat().st_size // 1024,
    )


# ============================================================
# HTML → Markdown
# ============================================================

def _convert_html(path: Path, out_dir: Path) -> ConvertResult:
    """使用 BeautifulSoup 将 HTML 转为 Markdown"""
    from bs4 import BeautifulSoup

    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")

    # 移除 script/style 标签
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    md_parts: list[str] = []

    # 标题
    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else path.stem
    md_parts.append(f"# {title}\n\n")
    md_parts.append(f"> 来源: {path.name}\n\n")

    # 提取正文
    body = soup.find("body") or soup

    # 遍历标签转为 Markdown
    _html_to_md(body, md_parts)

    # 统计段落
    para_count = len([p for p in body.find_all("p") if p.get_text(strip=True)])

    out_path = out_dir / f"{path.stem}.md"
    _write_md(out_path, "".join(md_parts))

    return ConvertResult(
        input_path=path,
        output_path=out_path,
        format="HTML",
        pages_or_chapters=para_count,
        size_kb=out_path.stat().st_size // 1024,
    )


def _html_to_md(element, result: list[str], list_depth: int = 0):
    """递归将 HTML 元素转为 Markdown"""
    from bs4 import NavigableString, Tag

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                result.append(f"{text} ")
            continue

        if not isinstance(child, Tag):
            continue

        tag_name = child.name.lower() if child.name else ""

        if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag_name[1])
            text = child.get_text(strip=True)
            if text:
                result.append(f"\n\n{'#' * level} {text}\n\n")

        elif tag_name == "p":
            text = child.get_text(strip=True)
            if text:
                result.append(f"{text}\n\n")

        elif tag_name in ("li",):
            indent = "  " * list_depth
            text = child.get_text(strip=True)
            if text:
                result.append(f"{indent}- {text}\n")

        elif tag_name in ("strong", "b"):
            result.append(f"**{child.get_text(strip=True)}**")

        elif tag_name in ("em", "i"):
            result.append(f"*{child.get_text(strip=True)}*")

        elif tag_name == "a":
            href = child.get("href", "")
            text = child.get_text(strip=True)
            if href and text:
                result.append(f"[{text}]({href})")
            elif text:
                result.append(text)

        elif tag_name == "img":
            alt = child.get("alt", "")
            src = child.get("src", "")
            if src:
                result.append(f"\n\n![{alt}]({src})\n\n")

        elif tag_name in ("blockquote",):
            text = child.get_text(strip=True)
            if text:
                result.append(f"\n> {text}\n\n")

        elif tag_name in ("pre", "code"):
            text = child.get_text()
            if text.strip():
                result.append(f"\n```\n{text}\n```\n\n")

        elif tag_name in ("br",):
            result.append("\n")

        elif tag_name in ("hr",):
            result.append("\n---\n")

        else:
            # 递归处理子元素
            _html_to_md(child, result, list_depth)


# ============================================================
# 辅助
# ============================================================

def _write_md(path: Path, content: str) -> None:
    """写入 Markdown 文件，自动处理重复情况"""
    base = path
    counter = 1
    while base.exists():
        base = path.parent / f"{path.stem}_{counter}.md"
        counter += 1
    base.write_text(content, encoding="utf-8")
